import pandas as pd
from docx import Document

# تحميل ملف Excel
excel_path = r'C:\Users\ramon\Desktop\recipes.xlsx'  # تحديث المسار حسب الحاجة
excel_file = pd.ExcelFile(excel_path)

# إنشاء مستند Word جديد
doc = Document()
doc.add_heading('Excel to Word Conversion', level=1)

# معالجة كل ورقة عمل في ملف Excel
for sheet_name in excel_file.sheet_names:
    df = pd.read_excel(excel_path, sheet_name=sheet_name)
    doc.add_heading(sheet_name, level=2)
    table = doc.add_table(rows=1, cols=len(df.columns))
    
    # إضافة رؤوس الأعمدة
    hdr_cells = table.rows[0].cells
    for i, column_name in enumerate(df.columns):
        hdr_cells[i].text = column_name
    
    # إضافة الصفوف
    for index, row in df.iterrows():
        row_cells = table.add_row().cells
        for i, value in enumerate(row):
            row_cells[i].text = str(value)

# حفظ مستند Word
word_path = r'C:\Users\ramon\Desktop\recipes.docx'
doc.save(word_path)

print(f"Converted Excel to Word and saved to {word_path}")
